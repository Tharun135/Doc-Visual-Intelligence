from pathlib import Path
from io import BytesIO
import time
import zipfile

from docx import Document
from pypdf import PdfReader

MAX_PDF_PAGES = 500
MAX_PARSE_SECONDS = 30
MAX_DOCX_UNCOMPRESSED_BYTES = 100 * 1024 * 1024


def _read_plain_text(filepath):
    with open(filepath, "r", encoding="utf-8") as file:
        text = file.read()
    # Normalize line endings and preserve structure
    return text.strip()


def _read_pdf_text(filepath):
    reader = PdfReader(filepath)
    if len(reader.pages) > MAX_PDF_PAGES:
        raise ValueError(f"PDF page count exceeds limit ({MAX_PDF_PAGES}).")

    started = time.monotonic()
    page_text = []
    for page in reader.pages:
        if time.monotonic() - started > MAX_PARSE_SECONDS:
            raise TimeoutError(f"PDF parsing exceeded {MAX_PARSE_SECONDS} seconds.")
        text = page.extract_text() or ""
        # Preserve structure by keeping non-empty pages with double newline separation
        if text.strip():
            page_text.append(text.strip())
    return "\n\n".join(page_text)


def _read_docx_text(filepath):
    with zipfile.ZipFile(filepath) as zf:
        uncompressed_total = sum(info.file_size for info in zf.infolist())
        if uncompressed_total > MAX_DOCX_UNCOMPRESSED_BYTES:
            raise ValueError("DOCX uncompressed size exceeds allowed limit.")

    started = time.monotonic()
    document = Document(filepath)
    text_parts = []
    for paragraph in document.paragraphs:
        if time.monotonic() - started > MAX_PARSE_SECONDS:
            raise TimeoutError(f"DOCX parsing exceeded {MAX_PARSE_SECONDS} seconds.")
        text = paragraph.text.strip()
        if text:
            text_parts.append(text)
    # Use double newlines to preserve section structure
    return "\n\n".join(text_parts)


def _read_plain_text_stream(file_stream):
    raw = file_stream.read()
    if isinstance(raw, bytes):
        text = raw.decode("utf-8", errors="replace")
    else:
        text = str(raw)
    return text.strip()


def _read_pdf_stream(file_stream):
    file_stream.seek(0)
    reader = PdfReader(file_stream)
    if len(reader.pages) > MAX_PDF_PAGES:
        raise ValueError(f"PDF page count exceeds limit ({MAX_PDF_PAGES}).")

    started = time.monotonic()
    page_text = []
    for page in reader.pages:
        if time.monotonic() - started > MAX_PARSE_SECONDS:
            raise TimeoutError(f"PDF parsing exceeded {MAX_PARSE_SECONDS} seconds.")
        text = page.extract_text() or ""
        if text.strip():
            page_text.append(text.strip())
    return "\n\n".join(page_text)


def _read_docx_stream(file_stream):
    file_stream.seek(0)
    raw = file_stream.read()
    with zipfile.ZipFile(BytesIO(raw)) as zf:
        uncompressed_total = sum(info.file_size for info in zf.infolist())
        if uncompressed_total > MAX_DOCX_UNCOMPRESSED_BYTES:
            raise ValueError("DOCX uncompressed size exceeds allowed limit.")

    started = time.monotonic()
    document = Document(BytesIO(raw))
    text_parts = []
    for paragraph in document.paragraphs:
        if time.monotonic() - started > MAX_PARSE_SECONDS:
            raise TimeoutError(f"DOCX parsing exceeded {MAX_PARSE_SECONDS} seconds.")
        text = paragraph.text.strip()
        if text:
            text_parts.append(text)
    return "\n\n".join(text_parts)


def extract_text(filepath):
    extension = Path(filepath).suffix.lower()

    if extension in {".txt", ".md", ".json"}:
        return _read_plain_text(filepath)
    if extension == ".pdf":
        return _read_pdf_text(filepath)
    if extension == ".docx":
        return _read_docx_text(filepath)

    raise ValueError(f"Unsupported file type: {extension}")


def extract_text_from_upload(file_storage):
    """Extract text from a Flask upload stream without writing persistent files."""
    filename = getattr(file_storage, "filename", "") or ""
    extension = Path(filename).suffix.lower()
    stream = file_storage.stream

    if extension in {".txt", ".md", ".json"}:
        stream.seek(0)
        return _read_plain_text_stream(stream)
    if extension == ".pdf":
        # PdfReader supports file-like streams directly.
        return _read_pdf_stream(stream)
    if extension == ".docx":
        return _read_docx_stream(stream)

    raise ValueError(f"Unsupported file type: {extension}")